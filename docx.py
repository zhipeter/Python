from docx import Document
from docx.shared import Inches
docu=Document()
docu.add_heading('The Head of Zhi',0)
p=docu.add_paragraph('郅朋 土木5班 1450856')
p.add_run('郅朋').bold=True
p.add_run('土木5班')
p.add_run('1450856').italic=True

docu.add_heading('Before')
docu.add_paragraph('I am a young boy, and pure',style='IntenseQuote')
docu.add_paragraph('Of course, I must change',style='ListBullet')
docu.add_picture('python.gif')

docu.add_page_break()

docu.save('zhi.docx')
