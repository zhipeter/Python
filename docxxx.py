from docx import Document
from docx.shared import Inches
f=open('zhi.docx','rb')
docu=Document(f)
docu.add_heading('The Head of Zhi',0)
p=docu.add_paragraph('郅朋 土木5班 1450856')
p.add_run('郅朋').bold=True
p.add_run('土木5班')
p.add_run('1450856').italic=True

docu.add_heading('Beforeghjjjjgjgjgjghjh')

docu.add_page_break()

f.close()