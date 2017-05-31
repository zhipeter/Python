from docx import Document
from wxpy import *

# bot=Bot(cache_path=True,console_qr=True)
bot=Bot(cache_path=True)

# gan=bot.friends().search('甘兆焯')[0]
# bing=bot.mps().search('小冰')[0]

@bot.register()
def print_others(msg):
	# with open('\home\ubuntu\Documents\log.txt','a') as f:
	# 	f.write(str(msg.receive_time)+'\n'+str(msg)+'\r\n')
	document=Document('zhi.docx')
	document.add_heading('The Head of Zhi',0)
	p=document.add_paragraph(str(msg.sender),style='List Bullet')
	if msg.type=='Text':
		text=document.add_paragraph(msg.text)
	elif msg.type=='Picture':
		file_name=msg.file_name
		img_path="E:\\E-Hen\\"+file_name
		msg.get_file(img_path)
		picture=document.add_picture(img_path，width=Inches(1.25))
	else:
		document.add_paragraph(str(msg))
		print("33")
	document.add_paragraph(str(msg.receive_time),style='Intense Quote')
	document.save('zhi.docx')

embed()